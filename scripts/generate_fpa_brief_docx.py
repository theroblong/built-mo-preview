#!/usr/bin/env python3
"""Generate docs/mo_fpa_team_brief.docx — native Word version of the HTML FP&A brief."""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'docs', 'mo_fpa_team_brief.docx')

# ── Palette ────────────────────────────────────────────────────────────────────
BLACK  = RGBColor(26, 26, 26)
DKGRAY = RGBColor(71, 85, 105)
GRAY   = RGBColor(100, 116, 139)
LGRAY  = RGBColor(148, 163, 184)
GREEN  = RGBColor(21, 128, 61)
AMBER  = RGBColor(146, 64, 14)
PURPLE = RGBColor(124, 58, 237)
WHITE  = RGBColor(255, 255, 255)

def _rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

# ── XML primitives ─────────────────────────────────────────────────────────────
def _cell_shade(cell, fill):
    fill = fill.lstrip('#')
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for s in tcPr.findall(qn('w:shd')): tcPr.remove(s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def _cell_borders(cell, *, top=None, right=None, bottom=None, left=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for b in tcPr.findall(qn('w:tcBorders')): tcPr.remove(b)
    tcB = OxmlElement('w:tcBorders')
    for name, spec in [('top', top), ('right', right), ('bottom', bottom), ('left', left)]:
        if spec is None: continue
        el = OxmlElement(f'w:{name}')
        el.set(qn('w:val'),   spec.get('val', 'single'))
        el.set(qn('w:sz'),    str(spec.get('sz', 4)))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), spec.get('color', '000000').lstrip('#'))
        tcB.append(el)
    tcPr.append(tcB)

def _no_tbl_borders(tbl):
    t = tbl._tbl
    tP = t.find(qn('w:tblPr'))
    if tP is None:
        tP = OxmlElement('w:tblPr'); t.insert(0, tP)
    for b in tP.findall(qn('w:tblBorders')): tP.remove(b)
    tB = OxmlElement('w:tblBorders')
    for e in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{e}')
        el.set(qn('w:val'), 'none'); el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), 'auto')
        tB.append(el)
    tP.append(tB)

def _para_shade(p, fill):
    fill = fill.lstrip('#')
    pPr = p._p.get_or_add_pPr()
    for s in pPr.findall(qn('w:shd')): pPr.remove(s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def _bottom_rule(p, color='1A1A1A', sz=6):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), str(sz))
    b.set(qn('w:space'), '4');    b.set(qn('w:color'), color.lstrip('#'))
    pBdr.append(b)
    pPr.append(pBdr)

def _force_font(r, name='Calibri'):
    r.font.name = name
    rPr = r._r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr'); r._r.insert(0, rPr)
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.insert(0, rf)
    rf.set(qn('w:ascii'), name); rf.set(qn('w:hAnsi'), name)

# ── Run / para helpers ─────────────────────────────────────────────────────────
def _r(para, text, *, bold=False, italic=False, size=10, color=BLACK):
    r = para.add_run(text)
    r.font.bold = bold; r.font.italic = italic
    r.font.size = Pt(size); r.font.color.rgb = color
    _force_font(r)
    return r

def _p(doc, *, sb=0, sa=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    return p

def _gap(doc, pts=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(); r.font.size = Pt(pts)
    return p

# ── Compound elements ──────────────────────────────────────────────────────────
def section_label(doc, text, txt_color, bg_hex):
    p = _p(doc, sb=18, sa=3)
    _para_shade(p, bg_hex)
    _r(p, f'  {text.upper()}  ', bold=True, size=8, color=txt_color)

def h2(doc, text):
    p = _p(doc, sb=4, sa=10)
    _r(p, text, bold=True, size=15, color=BLACK)
    _bottom_rule(p)

def h3(doc, text):
    p = _p(doc, sb=14, sa=6)
    _r(p, text.upper(), bold=True, size=9, color=DKGRAY)

def lead(doc, text):
    p = _p(doc, sa=10)
    _r(p, text, size=11, color=BLACK)

def body(doc, text, *, size=10, color=None, italic=False, bold=False, sa=8):
    color = color or BLACK
    p = _p(doc, sa=sa)
    _r(p, text, size=size, color=color, italic=italic, bold=bold)

def mixed(doc, parts, *, size=10, sa=8):
    """parts: list of (text, bold, italic, color)"""
    p = _p(doc, sa=sa)
    for text, b, i, c in parts:
        _r(p, text, bold=b, italic=i, size=size, color=c)

# ── Callout box (single-cell table with left accent border) ───────────────────
def callout(doc, paragraphs, *, bg='F8FAFC', accent='CBD5E1', title=None):
    """
    paragraphs: list where each item is either:
      - str  → plain text paragraph
      - list of (text, bold, italic, color) tuples → mixed-format paragraph
    title: optional bold prefix on the first paragraph
    """
    tbl = doc.add_table(rows=1, cols=1)
    _no_tbl_borders(tbl)
    cell = tbl.rows[0].cells[0]
    _cell_shade(cell, bg)
    _cell_borders(cell,
        top    = {'val': 'single', 'sz': 4,  'color': accent},
        right  = {'val': 'single', 'sz': 4,  'color': accent},
        bottom = {'val': 'single', 'sz': 4,  'color': accent},
        left   = {'val': 'single', 'sz': 18, 'color': accent},
    )

    def _fill(p_obj, content):
        p_obj.paragraph_format.left_indent  = Pt(8)
        p_obj.paragraph_format.right_indent = Pt(8)
        if isinstance(content, str):
            _r(p_obj, content, size=10, color=BLACK)
        else:
            for text, b, i, c in content:
                _r(p_obj, text, bold=b, italic=i, size=10, color=c)

    first_p = cell.paragraphs[0]
    first_p.paragraph_format.space_before = Pt(8)
    first_p.paragraph_format.space_after  = Pt(4)
    if title:
        _r(first_p, title + ' ', bold=True, size=10, color=BLACK)

    for idx, content in enumerate(paragraphs):
        if idx == 0:
            _fill(first_p, content)
        else:
            np = cell.add_paragraph()
            np.paragraph_format.space_before = Pt(2)
            np.paragraph_format.space_after  = Pt(4)
            _fill(np, content)

    pad = cell.add_paragraph()
    pad.paragraph_format.space_before = Pt(0)
    pad.paragraph_format.space_after  = Pt(6)
    _gap(doc, 8)

# ── Standard data table ────────────────────────────────────────────────────────
def std_table(doc, headers, rows, col_widths=None):
    """
    Each cell value can be:
      str                    → plain text
      (str, bool)            → (text, bold)
      (str, bool, bool, RGBColor) → (text, bold, italic, color)
    """
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'

    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hrow.cells[i]
        _cell_shade(c, '1A1A1A')
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent  = Pt(4)
        _r(p, h, bold=True, size=9, color=WHITE)

    for ri, row_data in enumerate(rows):
        bg = 'F9FAFB' if ri % 2 == 1 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            c = tbl.rows[ri + 1].cells[ci]
            _cell_shade(c, bg)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent  = Pt(4)
            if isinstance(val, tuple):
                text = val[0]; b = val[1] if len(val) > 1 else False
                i_   = val[2] if len(val) > 2 else False
                col  = val[3] if len(val) > 3 else BLACK
                _r(p, text, bold=b, italic=i_, size=9, color=col)
            else:
                _r(p, val, size=9, color=BLACK)

    if col_widths:
        for ci, w in enumerate(col_widths):
            for cell in tbl.columns[ci].cells:
                cell.width = Inches(w)

    _gap(doc, 10)

# ── 2-column card grid ─────────────────────────────────────────────────────────
def cards(doc, items, cols=2):
    """items: [(title, body_text), ...]"""
    W = 6.54  # content width in inches
    for offset in range(0, len(items), cols):
        batch = items[offset:offset + cols]
        tbl = doc.add_table(rows=1, cols=len(batch))
        _no_tbl_borders(tbl)
        col_w = W / cols
        for i, (title, body_text) in enumerate(batch):
            c = tbl.rows[0].cells[i]
            c.width = Inches(col_w)
            _cell_shade(c, 'F8FAFC')
            _cell_borders(c,
                top   = {'val': 'single', 'sz': 12, 'color': '1A1A1A'},
                right = {'val': 'none',   'sz': 0,  'color': 'auto'},
                bottom= {'val': 'none',   'sz': 0,  'color': 'auto'},
                left  = {'val': 'none',   'sz': 0,  'color': 'auto'},
            )
            pt = c.paragraphs[0]
            pt.paragraph_format.space_before = Pt(8); pt.paragraph_format.space_after = Pt(4)
            pt.paragraph_format.left_indent  = Pt(6); pt.paragraph_format.right_indent = Pt(8)
            _r(pt, title, bold=True, size=10, color=BLACK)
            pb = c.add_paragraph()
            pb.paragraph_format.space_before = Pt(0); pb.paragraph_format.space_after = Pt(8)
            pb.paragraph_format.left_indent  = Pt(6); pb.paragraph_format.right_indent = Pt(8)
            _r(pb, body_text, size=9, color=DKGRAY)
        _gap(doc, 6)
    _gap(doc, 8)

# ── Main build ─────────────────────────────────────────────────────────────────
def build():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    # ── COVER ─────────────────────────────────────────────────────────────────
    p = _p(doc, sa=14)
    _r(p, 'BUILT Bar  ·  Aevah Mo Demand Intelligence  ·  July 2026  ·  Confidential',
       size=9, color=GRAY, italic=True)

    p = _p(doc, sa=10)
    _r(p, 'Demand Intelligence Built for BUILT Finance', bold=True, size=24, color=BLACK)

    p = _p(doc, sa=10)
    _r(p, ('A comprehensive view of Mo’s forecasting, cannibalization detection, price elasticity '
           'analysis, causal insights, and explainability — built to complement and supercharge your '
           'existing FP&A process with continuously improving, account-level precision grounded in live SPINS data.'),
       size=11, color=BLACK)

    p = _p(doc, sa=20)
    _r(p, 'Prepared for: BUILT Finance & FP&A Leadership  ·  Aevah + BUILT  ·  July 17, 2026',
       size=9, color=GRAY)

    div = _p(doc, sa=16)
    _bottom_rule(div, sz=8)

    # ── KPI STRIP ─────────────────────────────────────────────────────────────
    kpi_data = [
        ('SPINS Data Alone',        '13.1%',
         'Weighted forecast error achievable from SPINS POS data — 5-quarter rolling average',
         'FFFFFF', '555555'),
        ("BUILT's Current Process",  '7–10%',
         'Corporate-level accuracy — a strong baseline aggregating across accounts and SKUs',
         'FFFBEB', '92400E'),
        ('Mo Unified Intelligence',  '4.4%',
         'Aevah unifies all data sources with blended machine learning — best Q1 2026 result',
         'F0FDF4', '15803D'),
        ('Elasticity Validation',    '63%',
         'Direction accuracy on clean, non-confounded price events',
         'F0FDF4', '15803D'),
        ('Value Case',              '$9–22M+',
         'Annual combined value across three workstreams — see breakdown below',
         'F0FDF4', '15803D'),
    ]
    kpi_tbl = doc.add_table(rows=1, cols=5)
    _no_tbl_borders(kpi_tbl)
    for i, (tag, num, lbl, bg, tx) in enumerate(kpi_data):
        c = kpi_tbl.rows[0].cells[i]
        c.width = Inches(1.31)
        _cell_shade(c, bg)
        _cell_borders(c,
            top   = {'val': 'single', 'sz': 4, 'color': 'E5E7EB'},
            right = {'val': 'single', 'sz': 4, 'color': 'E5E7EB'},
            bottom= {'val': 'single', 'sz': 4, 'color': 'E5E7EB'},
            left  = {'val': 'single', 'sz': 4, 'color': 'E5E7EB'},
        )
        tc = _rgb(tx)
        pt = c.paragraphs[0]
        pt.paragraph_format.space_before = Pt(8); pt.paragraph_format.space_after = Pt(2)
        pt.paragraph_format.left_indent  = Pt(4)
        _r(pt, tag.upper(), bold=True, size=7, color=tc)
        pn = c.add_paragraph()
        pn.paragraph_format.space_before = Pt(0); pn.paragraph_format.space_after = Pt(2)
        pn.paragraph_format.left_indent  = Pt(4)
        _r(pn, num, bold=True, size=18, color=tc)
        pl = c.add_paragraph()
        pl.paragraph_format.space_before = Pt(0); pl.paragraph_format.space_after = Pt(8)
        pl.paragraph_format.left_indent  = Pt(4); pl.paragraph_format.right_indent = Pt(4)
        _r(pl, lbl, size=7, color=GRAY)
    _gap(doc, 14)

    # ── VALUE BREAKDOWN ────────────────────────────────────────────────────────
    h3(doc, 'Value Estimate: How the $9–22M Is Built')
    std_table(doc,
        headers=['Workstream', 'Basis', 'Annual Estimate'],
        rows=[
            [('Forecast accuracy improvement', True),
             'BUILT baseline 7–10% → Mo 4.4% = 2.6–5.6pp × $1M/pp industry benchmark',
             '$2.6M–$5.6M'],
            [('Trade spend optimization', True),
             'Elasticity-guided reallocation from low-elasticity (Walmart ε=−0.25) to high-elasticity accounts (Ahold ε=−1.26); estimated on BUILT’s trade spend base',
             '$3M–$8M'],
            [('Cannibalization avoidance', True),
             'Early detection prevents over-investment in launches that cannibalize existing products; estimated on BUILT’s 3–5 significant launch events per year',
             '$3M–$8M'],
            [('Combined (current signal set)', True),
             ('Conservative to midpoint range; upper bound as promo calendar and additional inputs are integrated', False, True, DKGRAY),
             '$8.6M–$21.6M'],
        ],
        col_widths=[1.8, 3.5, 1.2],
    )
    body(doc,
         '$1M per 1pp is a widely cited CPG industry benchmark based on the relationship between '
         'forecast error and working capital, expedite costs, and markdown recovery at revenue scales '
         'similar to BUILT’s. Component estimates are directional; Aevah can model BUILT-specific '
         'values once current trade spend and launch budgets are shared.',
         size=8, color=DKGRAY, sa=12)

    # ── CAPACITY CALLOUT ───────────────────────────────────────────────────────
    callout(doc, title='Analyst Capacity: What Can Actually Be Done Today',
        paragraphs=[
            [('Brian’s analytics team runs approximately ', False, False, BLACK),
             ('40 analyst hours', True, False, BLACK),
             (' per four-week cycle. FP&A contributes another ', False, False, BLACK),
             ('40 hours', True, False, BLACK),
             ('. That is 80 combined hours available for demand intelligence work each cycle.', False, False, BLACK)],
            [('To recreate Mo’s core analysis workstreams manually — per-SKU × retailer '
              'forecasting, cannibalization scoring across the portfolio, and elasticity validation '
              'against price events — would require an estimated ', False, False, BLACK),
             ('250–535 additional analyst hours per cycle', True, False, BLACK),
             ('. That analysis is structurally out of reach without a significant headcount expansion.', False, False, BLACK)],
            [('Mo does not replace what your team does today. It delivers analysis that your team cannot '
              'do at current scale — which is exactly where compounding value lives.',
              False, True, _rgb('444444'))],
        ],
        bg='F9FAFB', accent='1A1A1A',
    )

    # ── SECTION 1: FORECASTING FOUNDATION ─────────────────────────────────────
    section_label(doc, 'Forecasting Foundation', BLACK, 'E2E8F0')
    h2(doc, 'The Forecasting Foundation')
    lead(doc,
        'SPINS Point-of-Sale data alone — using statistical methods — can get weighted forecast '
        'error (wMAPE) to around 13.1% across the portfolio. BUILT’s current process, aggregating '
        'data across accounts and SKUs, achieves a strong 7–10% at the corporate level. Mo goes '
        'further: by unifying all available data sources with a blended machine learning ensemble, Aevah '
        'reached 4.4% wMAPE in Q1 2026 on a 13-week hold-out across 164 active series. The goal is to '
        'continuously extend that precision down to individual SKU × retailer detail — '
        'compounding accuracy with every four-week retraining cycle.')

    callout(doc, title='What continuous retraining delivers:',
        paragraphs=[
            'Every four-week SPINS cycle, Mo retrains on the latest data — so accuracy reflects '
            'the current portfolio, not a model built six months ago. In testing, a stale model degraded '
            'nearly 18 percentage points versus a freshly retrained one. Your team gets a forecast signal '
            'that keeps pace with BUILT’s rapid SKU launches, retailer resets, and distribution '
            'changes — automatically, without additional analyst effort.',
        ],
        bg='F0FDF4', accent='16A34A',
    )

    h3(doc, '13-Week Retailer Forecast — Floor · Plan · Ceiling')
    body(doc,
         '[Chart: Figure 1 — Retailer-level 13-week forward projection (top 6 accounts by volume). '
         'Dark line = last known SPINS week. Shaded band = q10–q90 demand range. '
         'Dashed center = q50 operating plan. Each panel is a separate model series.]',
         size=9, color=LGRAY, italic=True)

    callout(doc, title='Supply chain payoff:',
        paragraphs=[
            'Three numbers per account, not one. Floor drives safety stock. Ceiling drives capacity '
            'reservation. Center drives the operating plan. When a single number is wrong, the entire '
            'inventory position is wrong. A calibrated range lets your team plan around realistic upside '
            'and downside simultaneously.',
        ],
        bg='F8FAFC', accent='CBD5E1',
    )

    h3(doc, 'A Continuously Improving Ensemble')
    body(doc,
        'What is an ensemble model? An ensemble combines multiple specialized models — each '
        'optimized for a different aspect of the forecasting problem — into one more accurate signal. '
        'Rather than relying on any single approach, the ensemble routes each forecast series through the '
        'methods that work best for that type of series and blends the results. The same principle applies '
        'in financial forecasting: combining multiple well-reasoned perspectives consistently outperforms '
        'any single forecast alone.',
        size=10, color=DKGRAY)
    body(doc,
        'Mo’s ensemble uses three layers working together: Light Gradient Boosting Machine (LightGBM) '
        'handles the core demand signal using 27+ SPINS-derived features. Bayesian Structural Time Series / '
        'CausalImpact (BSTS) isolates the causal effect of specific price and promotion events. SHapley '
        'Additive exPlanations (SHAP) translates every prediction into a ranked list of business drivers in '
        'plain English. The ensemble retrains every four weeks as new SPINS data arrives — matching '
        'your own planning cycle — so accuracy compounds rather than decays.',
        size=10, color=DKGRAY, sa=12)

    std_table(doc,
        headers=['Model Layer', 'What It Contributes', 'FP&A Benefit'],
        rows=[
            [('LightGBM (quantile)', True),
             'P10/P50/P90 demand forecast using distribution, velocity, price, promo, seasonality',
             'Calibrated floor/plan/ceiling for every SKU × retailer'],
            [('Bayesian Structural Time Series / CausalImpact (BSTS)', True),
             'Bayesian counterfactual: what would demand have been without a specific price or promo event?',
             'Validates trade spend ROI with causal evidence, not correlation'],
            [('SHAP (SHapley Additive exPlanations)', True),
             'Ranks the contribution of every feature to every forecast — in plain-English business terms',
             'Every prediction has an auditable paper trail your finance team can inspect'],
            [('4-week retraining', True),
             'Rolling window retrain each SPINS cycle; stale model degrades ~18pp without it',
             'Accuracy stays current as SKUs launch, retailers reset, and your portfolio grows'],
        ],
        col_widths=[2.0, 2.74, 1.8],
    )

    # ── SECTION 2: CANNIBALIZATION ─────────────────────────────────────────────
    section_label(doc, 'Cannibalization Intelligence', GREEN, 'F0FDF4')
    h2(doc, 'Cannibalization Detection')
    lead(doc,
        'BUILT expanded from roughly 3.5 SKUs to 7.5 SKUs in a year and is still growing. Every new '
        'launch creates the same FP&A question: are those units incremental demand, or are they '
        'transferring volume from an existing BUILT product? Answering it wrong leads to overstated '
        'forecasts, excess inventory commitments, and trade spend on launches that are eating your own '
        'portfolio. Mo tracks 8-week rolling SKU-to-SKU pressure signals across the entire portfolio — '
        'giving your team an early read before the miss shows up in a quarterly revenue call.')

    callout(doc, title='What Mo measures:',
        paragraphs=[
            '8-week Pearson correlation between co-stocked BUILT SKUs at the same retailer. A consistently '
            'negative co-movement signal — one SKU rising as another falls — is flagged as '
            'cannibalization pressure. Strength, direction, and confidence evolve weekly as the sales '
            'pattern matures.',
        ],
        bg='F0FDF4', accent='16A34A',
    )

    h3(doc, 'Where Cannibalization Shows Up in Your FP&A Work')
    std_table(doc,
        headers=['FP&A Scenario', 'Without Mo', 'With Mo'],
        rows=[
            [('New flavor launch forecast', True),
             'Estimate incremental units by analogy or comp SKU; cannibalization is assumed but not measured',
             'Mo tracks cross-SKU pressure from week 8 onward; signals whether new units are additive or transferred'],
            [('Pack size expansion (1ct → 4pk)', True),
             'Compare velocity manually; hard to separate channel mix from true demand substitution',
             'Mo controls for distribution (TDP) and price, isolating the demand transfer signal at the same retailer'],
            [('Revenue quality review', True),
             'Portfolio growth looks good in aggregate; internal share shifts are invisible until the quarter closes',
             'SKU-level portfolio pressure visible in real time; distinguishes earned growth from intra-brand substitution'],
            [('Inventory commitment for Q3 launch', True),
             'Commit to full build on optimistic incremental assumption; learn the transfer rate retrospectively',
             'Early cannibalization signal informs the commitment range before the inventory order is placed'],
        ],
        col_widths=[1.7, 2.42, 2.42],
    )

    callout(doc, title='So what?',
        paragraphs=[
            [('When the Q3 new SKU launch is in the operating plan, Mo gives your team the signal to '
              'distinguish additive demand from internal share shift — before you commit to inventory '
              'and trade spend. ', False, False, BLACK),
             ('Now what? ', True, False, BLACK),
             ('Share the launch calendar and Mo starts tracking cannibalization pressure from launch week 8 onward.',
              False, False, BLACK)],
        ],
        bg='F0FDF4', accent='16A34A',
    )

    # ── SECTION 3: PRICE ELASTICITY ────────────────────────────────────────────
    section_label(doc, 'Price Elasticity & Trade Spend', AMBER, 'FFFBEB')
    h2(doc, 'Price Elasticity & Trade Spend Intelligence')
    lead(doc,
        'Not every retailer responds to price the same way. Mo measures per-retailer price sensitivity '
        'from real SPINS transaction history — controlling for distribution, promotion, seasonality, '
        'and pack size lifecycle — so your team knows where a price move will move units and where '
        'it will not. This is the input trade teams have always needed for promo ROI decisions but have '
        'never had at account-level precision.')

    h3(doc, 'Per-Retailer Price Sensitivity (Confirmed Negative Elasticity)')
    std_table(doc,
        headers=['Retailer', 'Elasticity (ε)', 'Plain-English Read', 'Trade Spend Implication'],
        rows=[
            ['Walmart', '−0.245',
             'Inelastic — shoppers not highly price-sensitive at this account',
             'Small TPRs carry limited volume upside; spend may be better allocated elsewhere'],
            ['Kroger', '−0.590',
             'Moderate — price moves are felt; validated by December 2025 event',
             'Promo ROI calculations should use ε = −0.59; causal lift confirmed at +28.6% with full promo mechanics'],
            ['Ahold Delhaize', '−1.262',
             'Elastic — demand responds strongly to price changes',
             'Price moves require demand offset planning; strong ROI opportunity when promo conditions are right'],
            ['Whole Foods', '−0.445',
             'Muted sensitivity — consistent with premium positioning',
             'Price response is lower than mass/grocery; velocity and distribution are stronger levers here'],
        ],
        col_widths=[1.1, 0.9, 2.1, 2.44],
    )

    h3(doc, 'The Kroger Case — Causal Price Event Analysis')
    body(doc,
        'In December 2025, Brownie Batter 4pk at Kroger dropped from $10.99 to $10.14 (−7.7%). '
        'Mo’s BSTS counterfactual model isolated the price-only effect from everything else happening '
        'that period — broader seasonal uplift, display and feature support, distribution changes — '
        'and answered: how many units were genuinely above the counterfactual baseline because of the price '
        'change? Price alone explained +4.7% lift. The full promo package (price + display + feature) drove '
        '+28.6%. The gap is the promo mechanic value — quantified for the first time at event level.',
        size=10, color=DKGRAY, sa=8)
    body(doc,
         '[Chart: Figure 2 — Kroger BB4pk price event, December 2025. Top: focal vs. control retailer. '
         'Middle: Average Retail Price (ARP) trajectory. Bottom: TDP stable — rules out distribution as confound. '
         'Net causal lift: +28.6% above BSTS baseline.]',
         size=9, color=LGRAY, italic=True)

    callout(doc, title='Trade spend payoff:',
        paragraphs=[
            'You now know what the price-only contribution was (+4.7%) versus what the full promo package '
            'added (+28.6% net). The 23.7pp gap is the display and feature mechanic — and it can be '
            'evaluated against the trade spend cost to determine true ROI. This is the calculation your '
            'team has been doing manually, with far less precision.',
        ],
        bg='FFFBEB', accent='D97706',
    )

    h3(doc, 'Price Scenario Modeling — Forward-Looking')
    cards(doc, [
        ('What if we drop Kroger price by 5%?',
         'Mo uses ε = −0.59 for Kroger to estimate the demand response — and shows the confidence '
         'range around that estimate. Your trade team gets the expected unit lift and can compare it against the '
         'margin cost of the TPR before committing to the promotion.'),
        ('Where should we prioritize promo spend?',
         'Elastic accounts (Ahold ε = −1.26) respond strongly to price moves; inelastic accounts '
         '(Walmart ε = −0.25) do not. Mo gives you an account-level elasticity ranking so trade '
         'dollars go where they generate the most incremental volume.'),
    ])

    h3(doc, 'Future Capabilities Roadmap')
    body(doc,
        'The price elasticity foundation unlocks a set of adjacent capabilities currently in design. '
        'These build directly on top of the per-retailer, per-SKU elasticity estimates Mo already produces '
        '— no new data infrastructure required.')
    cards(doc, [
        ('Trade Spend Analytics',
         'Automated promo ROI analysis across every trade event, every four-week cycle. Compares actual lift '
         'against elasticity-predicted lift to identify over-funded and under-funded promotions in hindsight — '
         'and forward-projects ROI before committing spend.'),
        ('Assortment Intelligence',
         'Per-retailer ranking of which pack sizes and variants are gaining or losing velocity relative to the '
         'category. Surfaces format optimization opportunities before they show up in quarterly review.'),
        ('Promo Calendar Integration',
         'Forward-planned promotions as a direct model input. Known upcoming TPRs and display events feed into '
         'the forecast horizon, replacing manual conditioning on planned events after the baseline is generated.'),
        ('Velocity Benchmarking',
         'Per-store velocity for BUILT SKUs vs. the category average at each retailer. Answers “Is our '
         'velocity at Walmart good or not?” — the comparison that turns a number into an insight a '
         'trade team can act on.'),
    ])

    # ── SECTION 4: EXPLAINABILITY ──────────────────────────────────────────────
    section_label(doc, 'Explainability', PURPLE, 'FAF5FF')
    h2(doc, 'Explainability & Complete Audit Trail')
    lead(doc,
        'A finance team that has been burned before needs to know where every number came from. Mo provides '
        'a complete data lineage: from raw SPINS transactions, through every normalization and feature '
        'engineering step, to the model prediction — at any level of the hierarchy. Every number is '
        'inspectable. Every calculation is reproducible. Every forecast has a ranked list of business drivers '
        'that explains why it landed there.')

    h3(doc, 'SHAP Waterfall — Every Prediction Has a Business Reason')
    body(doc,
        'The SHAP waterfall shows exactly what moved the Brownie Batter 4pk forecast at Walmart. Each bar is '
        'a business factor with a direction and a magnitude. Demand momentum, TDP trajectory, and seasonal '
        'pattern account for the majority of the prediction. Actual: 27,317 units. Model: 28,472 units. '
        'Error: 4.2%. If a reviewer wants to understand why that number is what it is, they can read it '
        'directly from the chart — in business language, not model vocabulary.',
        size=10, color=DKGRAY, sa=8)
    body(doc,
         '[Chart: Figure 3 — SHAP waterfall: Brownie Batter 4pk at Walmart (Dec 2025 cutpoint). '
         'Base value = portfolio average demand. Each bar = one feature’s contribution to the final prediction.]',
         size=9, color=LGRAY, italic=True)

    h3(doc, 'Data Lineage & Provenance — Audit-Ready by Design')
    callout(doc, title='For due diligence, valuation, or board review:',
        paragraphs=[
            'Mo can document the operations performed on source data — what was kept, what was normalized, '
            'what was excluded and why — so that any reviewer can trace a forecast output all the way back '
            'to the raw SPINS transaction. This is the paper trail that normally requires an external firm and '
            'months of manual reconstruction. With Mo, it is built into every training run.',
        ],
        bg='FAF5FF', accent='7C3AED',
    )

    cards(doc, [
        ('Mo Chat — Ask in Plain English',
         '“Why is the Walmart forecast up 18% this quarter?” Mo responds with the driver decomposition '
         '— distribution change, velocity trend, seasonal pattern — grounded in live SPINS data. '
         'No spreadsheet required.'),
        ('Forecast Drawer',
         'Per-SKU, per-retailer floor/plan/ceiling for the next 13 weeks. Drill from portfolio to retailer '
         'to SKU in three clicks. Updated automatically every SPINS cycle.'),
        ('Auto-Generated Briefing',
         'Every four-week SPINS cycle regenerates this briefing automatically with updated accuracy metrics, '
         'updated forecasts, and updated elasticity. Your team sees the accuracy trend over time — '
         'not a one-time snapshot.'),
    ])

    # ── SECTION 5: GROWTH INTELLIGENCE ────────────────────────────────────────
    section_label(doc, 'Growth Intelligence', BLACK, 'E2E8F0')
    h2(doc, 'Growth Quality & New Product Visibility')
    lead(doc,
        'BUILT is not a steady-state brand. Distribution expansion can make total units rise even when '
        'per-store velocity is flat — and a simple trend forecast will mistake new doors for durable '
        'demand. Mo separates these two stories every week, so your team knows whether growth is earned '
        'or distribution-dependent before you plan the next quarter around it.')

    h3(doc, 'TDP Decomposition — Distribution vs. Velocity')
    body(doc,
        'At Walmart, Brownie Batter 4pk and Cookie Dough Chunk growth has been primarily a distribution '
        'story — Total Distribution Points (TDP, the number of stores carrying the product) rising '
        'while per-store velocity holds flat. Mo weights TDP features heavily in those series and knows '
        'not to project velocity-led acceleration that isn’t happening. When velocity does start '
        'running ahead of distribution, Mo sees it first — and the forecast reflects it.',
        size=10, color=DKGRAY, sa=8)
    body(doc,
         '[Chart: Figure 4 — Demand decomposition: BB4pk and CDC4pk at Walmart. '
         'Green = TDP index. Purple dashed = velocity index (sell-through per store). '
         'When TDP leads and velocity is flat, growth is a distribution story.]',
         size=9, color=LGRAY, italic=True)

    h3(doc, 'What Growth Decomposition Means for Your Plan')
    std_table(doc,
        headers=['Growth Type', 'What Mo Sees', 'FP&A Implication'],
        rows=[
            [('Distribution-led growth', True),
             'TDP rising; velocity stable or slightly declining (normalization)',
             'Growth is real but depends on continued shelf wins; plan conservatively if reset risk exists'],
            [('Velocity-led growth', True),
             'TDP stable or modest; velocity rising — consumers rebuying faster',
             'Higher-quality demand signal; more repeatable; less dependent on new retail relationships'],
            [('Promo-led volume', True),
             'Units spike during promo weeks; base demand is stable or declining',
             '~30% of BUILT portfolio is promo-linked; plan margin and inventory differently for this portion'],
            [('New SKU cannibalization', True),
             'New SKU velocity rising; adjacent BUILT SKU velocity declining at same retailer',
             'Total BUILT units may hold flat while mix shifts; revenue and margin quality change with it'],
        ],
        col_widths=[1.6, 2.5, 2.44],
    )

    callout(doc, title='Context that matters:',
        paragraphs=[
            'BUILT’s velocity is trending down slightly as distribution expands — but the brand '
            'remains the highest-velocity product in the top-10 bar set at key accounts. Mo holds this '
            'context in the model rather than flagging velocity softness as a negative signal when it is '
            'actually normalization.',
        ],
        bg='F8FAFC', accent='CBD5E1',
    )

    # ── SECTION 6: FINANCE OPERATIONS ─────────────────────────────────────────
    section_label(doc, 'Finance Operations', GREEN, 'F0FDF4')
    h2(doc, 'What Your Finance Team Gets')
    lead(doc,
        'Mo is a finance control layer for demand planning — not an analytics experiment. Your team '
        'gets tighter numbers, cleaner ROI reads, and a systematic way to prove out every figure in the '
        'operating plan, from raw SPINS data to the revenue call.')

    cards(doc, [
        ('Revenue Planning Gets a Tighter Number',
         '13-week retailer forecasts grounded in current SPINS — with a realistic range — replace '
         'single-point extrapolations that fail when a retailer resets or runs a promotion.'),
        ('Inventory Risk Gets a Range',
         'Floor/plan/ceiling per account gives supply chain a practical downside and upside. '
         'No more safety stock built on optimistic point forecasts.'),
        ('Trade Spend Gets Accountability',
         'Causal counterfactual shows what would have happened without the promo. Price-only lift vs. '
         'total event lift, separated and quantified, every time.'),
        ('New SKU Launches Get a Guard Rail',
         'Cannibalization pressure tracked from week 8 onward. Know whether a new flavor is growing the '
         'category or cannibalizing Brownie Batter before you commit to the full inventory build.'),
        ('Price Decisions Get Evidence',
         'Per-retailer elasticity tells you where price moves will lift units and where they won’t. '
         'Spend promo dollars where the math supports it.'),
        ('Analyst Capacity Gets Extended',
         'Brian’s analytics team and FP&A together have approximately 80 analyst hours per four-week '
         'cycle. Mo’s automation delivers the equivalent of 250–535 additional hours of analysis '
         'per cycle — structurally out of reach without headcount expansion. Analysts spend that freed '
         'time on interpretation, exception management, and stakeholder communication.'),
    ])

    h3(doc, 'Go Forward Plan: How Your Team Uses Mo')
    std_table(doc,
        headers=['Role', 'Screen / Tool', 'When They Use It', 'What They Walk Away With'],
        rows=[
            [('FP&A Analyst', True),
             'Forecast Accuracy Dashboard → SKU Retailer View',
             'First week of each 4-week SPINS cycle',
             'Updated wMAPE vs. prior cycle; which SKU × retailer pairs improved or degraded; ready for the revenue call without a manual SPINS pull'],
            [('Brian (Analytics Lead)', True),
             'Cannibalization Suite → Event Queue → Pre/Post Analysis',
             'When a new SKU launches or a quarterly review is due',
             'Early read on portfolio self-competition before it surfaces in the revenue call; recommendation on whether to adjust trade support or reposition a SKU'],
            [('Trade / Sales Finance', True),
             'Price Elasticity Suite → Scenario Modeling',
             'Before committing to a TPR or promotional event',
             'Per-account elasticity estimate with confidence range; expected unit lift and margin cost side by side; which accounts to prioritize and which to hold'],
            [('CFO / Finance Leadership', True),
             'Mo Chat → Board summary questions',
             'Board prep or quarterly close; ad hoc during revenue calls',
             'Plain-English answers with cited SPINS data; ranked driver explanation for any forecast miss; exportable summary for board package'],
            [('All Roles', True),
             'Mo Chat natural language interface',
             'Any time a question arises that would otherwise require a SPINS analyst',
             'Answer in seconds with a data citation, a confidence level, and a navigation link to the underlying screen for verification'],
        ],
        col_widths=[1.2, 1.7, 1.4, 2.24],
    )

    h3(doc, 'Inputs That Compound Accuracy Further')
    cards(doc, [
        ('H2 Promo Calendar',
         'Forward-looking promo dates convert retroactive inference to predictive modeling. '
         'Estimated +3–5pp wMAPE improvement on high-velocity, trade-heavy SKUs — equivalent to '
         '$3–5M at $1M per 1pp. One spreadsheet share from your trade planning team.'),
        ('Your Team’s Forecast Accuracy Baseline',
         'Sharing actual weighted forecast error on 3–5 key SKU × retailer pairs replaces the '
         'abstract benchmark with BUILT’s own number. This anchors the ROI claim in your actuals, '
         'not an industry estimate.'),
        ('Planned Distribution Changes',
         'Shelf resets, new retailer launches, planogram expansions as forward-looking TDP signals. '
         'Knowing a Walmart reset in advance closes the 4–8 week SPINS reflection lag entirely.'),
    ])

    # ── CLOSING ────────────────────────────────────────────────────────────────
    _gap(doc, 10)
    cls_tbl = doc.add_table(rows=1, cols=2)
    _no_tbl_borders(cls_tbl)
    for c in cls_tbl.rows[0].cells:
        _cell_shade(c, '0F172A')
        _cell_borders(c,
            top   = {'val': 'single', 'sz': 4, 'color': '334155'},
            right = {'val': 'none',   'sz': 0, 'color': 'auto'},
            bottom= {'val': 'none',   'sz': 0, 'color': 'auto'},
            left  = {'val': 'none',   'sz': 0, 'color': 'auto'},
        )

    def _dark_bullets(cell, heading, items):
        ph = cell.paragraphs[0]
        ph.paragraph_format.space_before = Pt(12); ph.paragraph_format.space_after = Pt(6)
        ph.paragraph_format.left_indent  = Pt(10)
        _r(ph, heading, bold=True, size=10, color=WHITE)
        for item in items:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent  = Pt(14)
            _r(p, '•  ' + item, size=9, color=WHITE)
        pad = cell.add_paragraph()
        pad.paragraph_format.space_before = Pt(0); pad.paragraph_format.space_after = Pt(10)

    _dark_bullets(cls_tbl.rows[0].cells[0], 'So What — Results Achieved', [
        '4.4% wMAPE at SKU×retailer level in Q1 2026 — tighter than BUILT’s corporate total, at a harder granularity',
        'Per-retailer price elasticity confirmed on real events: Kroger BB4pk +28.6% total lift, +4.7% price-only',
        'Cannibalization pressure tracked 8-week rolling across full BUILT portfolio',
        'Growth decomposed: distribution-led vs. velocity-led vs. promo-led — visible every week',
        'Every number has a ranked business driver explanation and a data lineage your team can audit',
    ])
    _dark_bullets(cls_tbl.rows[0].cells[1], 'Now What — The Next Steps', [
        'Share BUILT’s current forecast accuracy — makes the ROI comparison your own number, not an abstract benchmark',
        'Share the H2 promotional calendar — single highest-value input; estimated $3–5M accuracy gain',
        '90-day parallel run: every 4-week retrain generates an updated accuracy comparison vs. actuals',
        'By end of 90 days, the accuracy story belongs to BUILT — proven on your own data, your own SKUs, your own retailers',
    ])

    _gap(doc, 12)

    # ── FOOTER ─────────────────────────────────────────────────────────────────
    fp = _p(doc, sb=16, sa=0)
    _bottom_rule(fp, color='CBD5E1', sz=4)
    _r(fp,
       'Aevah Mo Demand Intelligence  ·  Powered by SPINS POS Data  ·  '
       'ML pipeline: MO_16–MO_49  ·  Generated July 17, 2026  ·  Confidential',
       size=8, color=LGRAY)

    doc.save(OUT)
    print(f'Saved: {OUT}')

if __name__ == '__main__':
    build()
